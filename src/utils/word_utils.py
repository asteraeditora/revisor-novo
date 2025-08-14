from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
import re
from copy import deepcopy

class WordDocumentHandler:
    """Classe para manipulação de documentos Word preservando formatação"""
    
    @staticmethod
    def read_document_complete(doc_path):
        """Lê documento preservando TODA estrutura, formatação e conteúdo"""
        doc = Document(doc_path)
        content = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Parágrafo
                para = Paragraph(element, doc)
                content.append({
                    'type': 'paragraph',
                    'text': para.text,
                    'style': para.style.name if para.style else None,
                    'alignment': para.alignment,
                    'runs': WordDocumentHandler._extract_runs(para),
                    'element': element
                })
            elif element.tag.endswith('tbl'):  # Tabela
                table = Table(element, doc)
                content.append({
                    'type': 'table',
                    'data': WordDocumentHandler._extract_table_data(table),
                    'element': element
                })
        
        return content
    
    @staticmethod
    def _extract_runs(paragraph):
        """Extrai runs com formatação detalhada"""
        runs_data = []
        for run in paragraph.runs:
            run_data = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size.pt if run.font.size else None,
                'color': None
            }
            
            # Verifica se é um link
            if WordDocumentHandler._is_hyperlink(run):
                run_data['is_link'] = True
            
            # Cor do texto
            if run.font.color and run.font.color.rgb:
                run_data['color'] = str(run.font.color.rgb)
            
            runs_data.append(run_data)
        
        return runs_data
    
    @staticmethod
    def _is_hyperlink(run):
        """Verifica se o run contém um hyperlink"""
        # Verificação simplificada por padrão de URL
        url_pattern = r'https?://[^\s]+'
        return bool(re.search(url_pattern, run.text))
    
    @staticmethod
    def _extract_table_data(table):
        """Extrai dados da tabela preservando formatação"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_content = []
                for paragraph in cell.paragraphs:
                    cell_content.append({
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else None,
                        'runs': WordDocumentHandler._extract_runs(paragraph)
                    })
                row_data.append(cell_content)
            table_data.append(row_data)
        return table_data
    
    @staticmethod
    def apply_correction_preserving_format(paragraph, error_text, correction_text):
        """Aplica correção preservando ABSOLUTAMENTE TODA a formatação"""
        
        # Verifica se o erro existe no parágrafo
        full_text = ""
        for run in paragraph.runs:
            full_text += run.text
        
        if error_text not in full_text:
            return False
        
        error_pos = full_text.find(error_text)
        error_end = error_pos + len(error_text)
        
        # ABORDAGEM SIMPLIFICADA: trabalha run por run
        current_pos = 0
        correction_applied = False
        
        for run in paragraph.runs:
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            # Se este run contém parte do erro
            if run_start < error_end and run_end > error_pos:
                # Calcula que parte do run deve ser mantida/alterada
                
                # Início do texto a manter (antes do erro)
                keep_start = ""
                if run_start < error_pos:
                    keep_start = run.text[:error_pos - run_start]
                
                # Fim do texto a manter (depois do erro)
                keep_end = ""
                if run_end > error_end:
                    keep_end = run.text[error_end - run_start:]
                
                # Parte da correção que vai neste run
                correction_part = ""
                if not correction_applied:
                    correction_part = correction_text
                    correction_applied = True
                
                # Atualiza APENAS o texto, mantendo TODA a formatação
                run.text = keep_start + correction_part + keep_end
                
                # Se o erro estava completamente neste run, podemos parar
                if run_start <= error_pos and run_end >= error_end:
                    return True
                
            elif run_start >= error_pos and run_end <= error_end:
                # Este run está completamente dentro do erro - limpa ele
                run.text = ""
            
            current_pos += len(run.text)
        
        return correction_applied
    
    @staticmethod
    def _apply_runs_to_paragraph(paragraph, original_runs, new_text):
        """Aplica formatação dos runs ao novo texto"""
        # Se não há runs ou texto vazio
        if not original_runs or not new_text:
            paragraph.add_run(new_text)
            return
        
        # Detecta e preserva links
        url_pattern = r'(https?://[^\s]+)'
        parts = re.split(url_pattern, new_text)
        
        current_pos = 0
        for part in parts:
            if not part:
                continue
                
            run = paragraph.add_run(part)
            
            # Se é URL, aplica formatação de link
            if re.match(url_pattern, part):
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
            else:
                # Aplica formatação do run original mais próximo
                if original_runs:
                    # Usa formatação do primeiro run como base
                    orig_run = original_runs[0]
                    if orig_run.get('bold'):
                        run.bold = True
                    if orig_run.get('italic'):
                        run.italic = True
                    if orig_run.get('underline') and not re.match(url_pattern, part):
                        run.underline = True
                    if orig_run.get('font_name'):
                        run.font.name = orig_run['font_name']
                    if orig_run.get('font_size'):
                        run.font.size = Pt(orig_run['font_size'])
                    if orig_run.get('color') and not re.match(url_pattern, part):
                        try:
                            color_hex = orig_run['color']
                            if isinstance(color_hex, str) and len(color_hex) == 6:
                                run.font.color.rgb = RGBColor(
                                    int(color_hex[0:2], 16),
                                    int(color_hex[2:4], 16),
                                    int(color_hex[4:6], 16)
                                )
                        except:
                            pass
    
    @staticmethod
    def extract_images_info(doc_path):
        """Extrai informações sobre imagens no documento"""
        doc = Document(doc_path)
        images_info = []
        
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.target_ref:
                images_info.append({
                    'index': i,
                    'relationship_id': rel.rId,
                    'target': rel.target_ref
                })
        
        return images_info