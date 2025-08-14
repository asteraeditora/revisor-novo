import json
import os
import re
import difflib
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def adicionar_paragrafo_com_destaque_por_palavra(doc, texto_original, texto_revisado):
    """
    Adiciona um parágrafo ao documento, destacando as diferenças PALAVRA POR PALAVRA.
    """
    p = doc.add_paragraph()
    p.add_run("Resultado Visual: ").bold = True
    
    palavras_orig = re.findall(r'\w+|\W+', texto_original)
    palavras_rev = re.findall(r'\w+|\W+', texto_revisado)

    matcher = difflib.SequenceMatcher(None, palavras_orig, palavras_rev)
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace':
            run_del = p.add_run("".join(palavras_orig[i1:i2]))
            run_del.font.strike = True
            run_del.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            
            run_add = p.add_run("".join(palavras_rev[j1:j2]))
            run_add.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            run_add.font.bold = True
            
        elif tag == 'delete':
            run_del = p.add_run("".join(palavras_orig[i1:i2]))
            run_del.font.strike = True
            run_del.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            
        elif tag == 'insert':
            run_add = p.add_run("".join(palavras_rev[j1:j2]))
            run_add.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            run_add.font.bold = True
            
        elif tag == 'equal':
            p.add_run("".join(palavras_orig[i1:i2]))

def adicionar_separador(doc):
    """Adiciona um parágrafo de espaçamento e uma linha horizontal."""
    doc.add_paragraph() # Espaço extra
    separator_p = doc.add_paragraph()
    run = separator_p.add_run('—' * 70)
    run.font.size = Pt(8)
    separator_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Espaço extra

def processar_formato_agrupado(data, doc):
    """Processa o JSON no formato original (agrupado por página)."""
    correcoes_por_pagina = data.get("correções_por_página", {})
    
    is_first_page_section = True
    for nome_pagina, dados_pagina in sorted(correcoes_por_pagina.items()):
        if not is_first_page_section:
            adicionar_separador(doc)
        is_first_page_section = False
        
        numero_pagina = nome_pagina.split('_')[-1]
        doc.add_heading(f'Correções na Página {numero_pagina}', level=1)
        
        for i, correcao in enumerate(dados_pagina.get("correções", [])):
            # ... (código do item de correção)
            p_titulo = doc.add_paragraph()
            p_titulo.add_run(f"Item de Correção {i+1} (Parágrafo: {correcao.get('paragrafo', 'N/A')})").bold = True
            adicionar_paragrafo_com_destaque_por_palavra(doc, correcao.get("erro", ""), correcao.get("correcao", ""))
            
            # Adiciona um separador mais simples entre itens da mesma página
            if i < len(dados_pagina.get("correções", [])) - 1:
                doc.add_paragraph('---' * 10)

def processar_formato_lista(data, doc):
    """Processa o JSON no novo formato (lista plana de 'todas_mudancas')."""
    mudancas_agrupadas = {}
    for mudanca in data.get("todas_mudancas", []):
        pagina = mudanca.get("page")
        if pagina is not None:
            if pagina not in mudancas_agrupadas:
                mudancas_agrupadas[pagina] = []
            mudancas_agrupadas[pagina].append(mudanca)

    is_first_page_section = True
    for pagina in sorted(mudancas_agrupadas.keys()):
        if not is_first_page_section:
            adicionar_separador(doc)
        is_first_page_section = False

        doc.add_heading(f'Correções na Página {pagina}', level=1)
        correcoes_da_pagina = mudancas_agrupadas[pagina]

        for i, correcao in enumerate(correcoes_da_pagina):
            # ... (código do item de correção)
            p_titulo = doc.add_paragraph()
            p_titulo.add_run(f"Item de Correção {i+1} (Parágrafo: {correcao.get('paragraph_number', 'N/A')})").bold = True
            adicionar_paragrafo_com_destaque_por_palavra(doc, correcao.get("original_full", ""), correcao.get("revised_full", ""))

            # Adiciona um separador mais simples entre itens da mesma página
            if i < len(correcoes_da_pagina) - 1:
                 doc.add_paragraph('---' * 10)

def criar_relatorio_de_arquivo(json_file):
    """Função principal que detecta o formato do JSON e gera o relatório."""
    print(f"--- Processando arquivo: {json_file} ---")
    nome_base = os.path.splitext(json_file)[0]
    output_docx = f'relatorio_{nome_base}.docx'

    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"ERRO: Não foi possível ler ou decodificar o arquivo '{json_file}'. Detalhes: {e}")
        return

    doc = Document()
    doc.add_heading('Relatório de Correções do Livro', level=0)
    doc.add_paragraph(f"Arquivo de origem: {json_file}\n")

    if 'correções_por_página' in data:
        print(f"Detectado formato 'Agrupado por Página'. Gerando relatório...")
        processar_formato_agrupado(data, doc)
    elif 'todas_mudancas' in data:
        print(f"Detectado formato 'Lista de Mudanças'. Gerando relatório...")
        processar_formato_lista(data, doc)
    else:
        print(f"AVISO: Estrutura de JSON não reconhecida no arquivo '{json_file}'.")
        return

    try:
        doc.save(output_docx)
        print(f"SUCESSO: Relatório '{output_docx}' foi criado.")
    except Exception as e:
        print(f"ERRO: Ocorreu um erro ao salvar o arquivo Word '{output_docx}'. Detalhes: {e}")

# --- INÍCIO DA EXECUÇÃO ---
if __name__ == "__main__":
    arquivos_json = [f for f in os.listdir('.') if f.endswith('.json')]

    if not arquivos_json:
        print("Nenhum arquivo .json foi encontrado na pasta.")
    else:
        print(f"Encontrados {len(arquivos_json)} arquivos JSON. Iniciando processamento...\n")
        for nome_do_arquivo in arquivos_json:
            criar_relatorio_de_arquivo(nome_do_arquivo)
            print("-" * 50)
    
    print("\nProcessamento de todos os arquivos concluído.")