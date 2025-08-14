import os
import shutil
import logging
import json
import re
from typing import List, Dict, Tuple
from docx import Document
from ..utils.word_utils import WordDocumentHandler
from ..utils.api_client import OpenAIClient

class DocumentProcessor:
    """Processa documentos Word - Versão COMPLETA que processa TUDO"""
    
    def __init__(self, api_key: str, model: str = "gpt-4.1"):
        self.api_client = OpenAIClient(api_key, model)
        self.api_key = api_key
        self.model = model
        self.word_handler = WordDocumentHandler()
        self.logger = logging.getLogger(__name__)
        self.max_chunk_size = 10000  # Aproveita janela do gpt-4.1
    
    def process_document(self, input_path: str, output_path: str, callback=None):
        """Processa documento corrigindo TODOS os erros reais"""
        try:
            self.logger.info("=== PROCESSAMENTO COMPLETO INICIADO ===")
            
            # 1. Copia o arquivo original
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            shutil.copy2(input_path, output_path)
            
            # 2. Abre ambos os documentos
            original_doc = Document(input_path)
            doc = Document(output_path)
            
            # 3. Mapeia TODOS os textos (parágrafos + tabelas)
            all_texts = []
            text_counter = 0
            
            # Rastreia contexto (para identificar citações/poemas)
            in_citation = False
            citation_start = -1
            in_poem = False
            poem_start = -1
            
            # Parágrafos normais
            self.logger.info("Mapeando parágrafos...")
            for i, (orig_para, para) in enumerate(zip(original_doc.paragraphs, doc.paragraphs)):
                text = para.text.strip()
                
                # Detecta início de citação/poema
                if any(marker in text.lower() for marker in [
                    'leia o texto:', 'observe o poema:', 'leia o poema:',
                    'texto:', 'poema:', 'leia a seguir:', 'leia:',
                    'o poema a seguir', 'o texto abaixo', 'a poesia'
                ]):
                    if 'poema' in text.lower() or 'poesia' in text.lower():
                        in_poem = True
                        poem_start = i + 1
                    else:
                        in_citation = True
                        citation_start = i + 1
                
                # Detecta fim de citação/poema (linha vazia ou nova questão)
                if (in_citation or in_poem) and (not text or re.match(r'^\d+[\.\)]', text)):
                    in_citation = False
                    in_poem = False
                
                if text:
                    text_counter += 1
                    
                    # Marca se está dentro de citação/poema
                    is_citation_content = (in_citation and i >= citation_start)
                    is_poem_content = (in_poem and i >= poem_start)
                    
                    # Detecta se é verso de poema pela estrutura
                    is_verse = (
                        is_poem_content or
                        (len(text) < 60 and 
                         not text.endswith('.') and 
                         not any(c in text for c in [',', ';']) and
                         i > 0 and len(doc.paragraphs[i-1].text.strip()) < 60)
                    )
                    
                    # Detecta se é realmente protegido
                    is_protected = (
                        self._is_really_protected(para.text) or 
                        is_citation_content or 
                        is_poem_content or
                        is_verse
                    )
                    
                    protection_reason = None
                    if is_poem_content or is_verse:
                        protection_reason = 'poema/verso'
                    elif is_citation_content:
                        protection_reason = 'citação'
                    elif self._is_really_protected(para.text):
                        protection_reason = 'outro'
                    
                    all_texts.append({
                        'index': text_counter,
                        'doc_index': i,
                        'original_text': orig_para.text,
                        'current_text': para.text,
                        'paragraph_obj': para,
                        'type': 'paragraph',
                        'protected': is_protected,
                        'protection_reason': protection_reason,
                        'location': f'Parágrafo {text_counter}'
                    })
            
            # Tabelas - SEMPRE processa
            self.logger.info("Mapeando tabelas...")
            for t_idx, (orig_table, table) in enumerate(zip(original_doc.tables, doc.tables)):
                for r_idx, (orig_row, row) in enumerate(zip(orig_table.rows, table.rows)):
                    for c_idx, (orig_cell, cell) in enumerate(zip(orig_row.cells, row.cells)):
                        for p_idx, (orig_para, para) in enumerate(zip(orig_cell.paragraphs, cell.paragraphs)):
                            if para.text.strip():
                                text_counter += 1
                                
                                # Tabelas não são automaticamente protegidas
                                is_protected = self._is_really_protected(para.text)
                                
                                all_texts.append({
                                    'index': text_counter,
                                    'doc_index': f"table_{t_idx}_{r_idx}_{c_idx}_{p_idx}",
                                    'original_text': orig_para.text,
                                    'current_text': para.text,
                                    'paragraph_obj': para,
                                    'type': 'table',
                                    'protected': is_protected,
                                    'location': f'Tabela {t_idx+1}, Célula ({r_idx+1},{c_idx+1})'
                                })
            
            self.logger.info(f"Total de textos mapeados: {len(all_texts)}")
            
            # 4. Separa processáveis
            processable = [t for t in all_texts if not t['protected']]
            protected = [t for t in all_texts if t['protected']]
            
            self.logger.info(f"Textos processáveis: {len(processable)}")
            self.logger.info(f"Textos protegidos: {len(protected)} (apenas alternativas e BNCC)")
            
            # 5. Cria blocos com todos os textos processáveis
            blocks = self._create_mixed_blocks(processable)
            self.logger.info(f"Criados {len(blocks)} blocos para análise")
            
            # 6. Processa cada bloco
            all_corrections = []
            corrections_by_index = {}  # Para rastreamento
            
            for block_idx, block in enumerate(blocks):
                if callback:
                    callback(block_idx + 1, len(blocks), 
                            f"Analisando bloco {block_idx+1}/{len(blocks)}")
                
                # Info do bloco
                block_info = f"Bloco {block_idx+1}: {len(block)} textos"
                if any(t['type'] == 'table' for t in block):
                    table_count = sum(1 for t in block if t['type'] == 'table')
                    block_info += f" ({table_count} de tabelas)"
                self.logger.info(block_info)
                
                # Prepara texto
                block_text = self._prepare_mixed_block(block)
                
                # Analisa
                corrections = self.api_client.identify_errors_precise(block_text, block_idx)
                
                if corrections:
                    self.logger.info(f"  {len(corrections)} correções sugeridas")
                    
                    for corr in corrections:
                        # Encontra o texto correspondente
                        text_data = self._find_text_in_block(block, corr)
                        
                        if text_data:
                            # Valida e aplica
                            if self._should_apply_correction(text_data, corr):
                                success = self._apply_correction_safe(text_data, corr)
                                
                                if success:
                                    correction_record = {
                                        'block': block_idx + 1,
                                        'text_index': text_data['index'],
                                        'location': text_data['location'],
                                        'type': text_data['type'],
                                        'error': corr.get('error', ''),
                                        'correction': corr.get('correction', ''),
                                        'error_type': corr.get('type', ''),
                                        'original_text': text_data['original_text'],
                                        'corrected_text': text_data['paragraph_obj'].text
                                    }
                                    all_corrections.append(correction_record)
                                    corrections_by_index[text_data['index']] = correction_record
                                    
                                    self.logger.info(f"  ✓ Aplicada em {text_data['location']}: '{corr['error']}' → '{corr['correction']}'")
            
            # 7. Verifica se há mudanças não detectadas pela API
            self.logger.info("Verificando integridade das correções...")
            for text_data in all_texts:
                current = text_data['paragraph_obj'].text
                original = text_data['original_text']
                
                if current != original and text_data['index'] not in corrections_by_index:
                    # Mudança não registrada!
                    self.logger.warning(f"Mudança não detectada no texto {text_data['index']}: {text_data['location']}")
                    
                    # Analisa a diferença
                    diff = self._analyze_difference(original, current)
                    all_corrections.append({
                        'block': 'auto',
                        'text_index': text_data['index'],
                        'location': text_data['location'],
                        'type': text_data['type'],
                        'error': diff['error'],
                        'correction': diff['correction'],
                        'error_type': 'auto-detectado',
                        'original_text': original,
                        'corrected_text': current
                    })
            
            # 8. Salva documento
            doc.save(output_path)
            
            self.logger.info(f"=== PROCESSAMENTO CONCLUÍDO ===")
            self.logger.info(f"Total de correções: {len(all_corrections)}")
            self.logger.info(f"Em parágrafos: {sum(1 for c in all_corrections if c['type'] == 'paragraph')}")
            self.logger.info(f"Em tabelas: {sum(1 for c in all_corrections if c['type'] == 'table')}")
            
            # 9. Salva relatório
            self._save_complete_report(output_path, all_corrections, protected)
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Erro no processamento: {str(e)}")
            raise
    
    def _is_really_protected(self, text: str) -> bool:
        """Protege APENAS o que realmente não deve ser alterado"""
        # Alternativas de questões (a), b), c) etc
        if re.match(r'^[a-eA-E][\)\.][\s]', text):
            return True
        
        # Códigos BNCC
        if re.search(r'\b(EF\d{2}[A-Z]{2}\d{2}|EM\d{2}[A-Z]{2}\d{2})\b', text):
            return True
        
        # Links completos (não corrige URLs)
        if re.search(r'https?://[^\s]+', text):
            return True
        
        # Referências bibliográficas (padrão ABNT)
        if re.match(r'^[A-Z]+,\s+[A-Z][a-z]+', text) and 'Editora' in text:
            return True
        
        # Gabarito
        if text.strip().upper().startswith('GABARITO'):
            return True
        
        # Citações diretas longas (mais de 3 linhas geralmente)
        if text.startswith('    ') or text.startswith('\t'):  # Indentado
            return True
        
        return False
    
    def _create_mixed_blocks(self, texts: List[Dict]) -> List[List[Dict]]:
        """Cria blocos misturando parágrafos e tabelas"""
        blocks = []
        current_block = []
        current_size = 0
        
        for text_data in texts:
            text_size = len(text_data['current_text'])
            
            # Cria novo bloco se necessário
            if current_size + text_size > self.max_chunk_size or len(current_block) >= 100:
                if current_block:
                    blocks.append(current_block)
                current_block = [text_data]
                current_size = text_size
            else:
                current_block.append(text_data)
                current_size += text_size
        
        if current_block:
            blocks.append(current_block)
        
        return blocks
    
    def _prepare_mixed_block(self, block: List[Dict]) -> str:
        """Prepara bloco com parágrafos e tabelas"""
        block_text = f"BLOCO COM {len(block)} TEXTOS\n\n"
        block_text += "Corrija TODOS os erros de português encontrados.\n\n"
        
        for text_data in block:
            block_text += f"[TEXTO {text_data['index']}]\n"
            block_text += f"[TIPO: {text_data['type'].upper()}]\n"
            
            if text_data['type'] == 'table':
                block_text += f"[LOCALIZAÇÃO: {text_data['location']}]\n"
            
            # Marca tipo de conteúdo
            content_type = self._detect_content_type(text_data['current_text'])
            block_text += f"[CONTEÚDO: {content_type}]\n"
            
            block_text += text_data['current_text'] + "\n"
            block_text += f"[FIM_TEXTO_{text_data['index']}]\n\n"
        
        return block_text
    
    def _detect_content_type(self, text: str) -> str:
        """Detecta o tipo de conteúdo com mais precisão"""
        # Título/Cabeçalho
        if len(text) < 80 and not text.endswith(('.', '!', '?', ':')):
            return "TÍTULO/CABEÇALHO"
        
        # Item de lista
        elif text.strip().startswith(('•', '-', '1.', '2.', 'a.', 'I.', 'II.')):
            return "ITEM DE LISTA"
        
        # Poema (versos curtos, sem ponto final)
        elif len(text) < 60 and not text.endswith('.') and not any(c in text for c in [',', ';']):
            return "POSSÍVEL VERSO"
        
        # URL/Link
        elif re.search(r'https?://', text):
            return "CONTÉM LINK"
        
        # Citação com aspas
        elif text.count('"') >= 2 or text.count('"') >= 2:
            return "CITAÇÃO"
        
        # Referência bibliográfica
        elif re.search(r'\b\d{4}\b', text) and any(word in text for word in ['Editora', 'ed.', 'p.', 'In:']):
            return "REFERÊNCIA"
        
        else:
            return "TEXTO NORMAL"
    
    def _find_text_in_block(self, block: List[Dict], correction: Dict) -> Dict:
        """Encontra texto no bloco pela correção"""
        # Tenta pelo número do texto
        text_num = correction.get('paragraph', correction.get('text', 0))
        error_text = correction.get('error', '')
        
        for text_data in block:
            if text_data['index'] == text_num:
                if error_text in text_data['current_text']:
                    return text_data
        
        # Tenta só pelo conteúdo
        for text_data in block:
            if error_text in text_data['current_text']:
                return text_data
        
        return None
    
    def _should_apply_correction(self, text_data: Dict, correction: Dict) -> bool:
        """Validação antes de aplicar - equilibrada"""
        text = text_data['current_text']
        error = correction.get('error', '')
        fix = correction.get('correction', '')
        
        # Verifica se o erro existe
        if error not in text:
            return False
        
        # Não altera demonstrativos
        if error.lower() in ['este', 'esse', 'esta', 'essa', 'isto', 'isso']:
            self.logger.info(f"Rejeitada: demonstrativo {error}")
            return False
        
        # Proteção especial para URLs
        if 'http' in text and 'http' in error:
            self.logger.info(f"Rejeitada: parte de URL")
            return False
        
        # Cuidado com citações - verifica se está dentro de aspas
        if any(quote in text for quote in ['"', '"', '"']):
            # Encontra posição do erro
            error_pos = text.find(error)
            if error_pos > -1:
                # Conta aspas antes do erro
                before_text = text[:error_pos]
                quotes_before = sum(before_text.count(q) for q in ['"', '"', '"'])
                
                # Se número ímpar, está dentro de aspas
                if quotes_before % 2 == 1:
                    self.logger.info(f"Rejeitada: dentro de citação")
                    return False
        
        # Se é poema/verso (identificado pelo tipo de conteúdo)
        if text_data.get('protection_reason') == 'citação/poema':
            self.logger.info(f"Rejeitada: conteúdo de poema/citação")
            return False
        
        return True
    
    def _apply_correction_safe(self, text_data: Dict, correction: Dict) -> bool:
        """Aplica correção com segurança"""
        try:
            paragraph = text_data['paragraph_obj']
            original_text = paragraph.text
            error = correction.get('error', '')
            fix = correction.get('correction', '')
            
            if not error or not fix or error not in original_text:
                return False
            
            # Aplica uma vez
            new_text = original_text.replace(error, fix, 1)
            
            # Validação final
            if abs(len(new_text) - len(original_text)) > 20:
                return False
            
            paragraph.text = new_text
            return True
            
        except Exception as e:
            self.logger.error(f"Erro ao aplicar: {str(e)}")
            return False
    
    def _analyze_difference(self, original: str, current: str) -> Dict:
        """Analisa diferença entre textos"""
        import difflib
        
        # Casos especiais
        if current == original + '.':
            return {'error': '[faltava ponto final]', 'correction': '.'}
        
        # Análise por palavras
        s = difflib.SequenceMatcher(None, original.split(), current.split())
        
        for tag, i1, i2, j1, j2 in s.get_opcodes():
            if tag == 'replace':
                return {
                    'error': ' '.join(original.split()[i1:i2]),
                    'correction': ' '.join(current.split()[j1:j2])
                }
            elif tag == 'insert':
                return {
                    'error': '[faltando]',
                    'correction': ' '.join(current.split()[j1:j2])
                }
            elif tag == 'delete':
                return {
                    'error': ' '.join(original.split()[i1:i2]),
                    'correction': '[removido]'
                }
        
        return {'error': 'mudança detectada', 'correction': 'texto alterado'}
    
    def _save_complete_report(self, output_path: str, corrections: List[Dict], protected: List[Dict]):
        """Salva relatório completo"""
        report_path = output_path.replace('.docx', '_relatorio_completo.json')
        
        # Estatísticas
        stats = {
            'total_correcoes': len(corrections),
            'em_paragrafos': sum(1 for c in corrections if c['type'] == 'paragraph'),
            'em_tabelas': sum(1 for c in corrections if c['type'] == 'table'),
            'auto_detectadas': sum(1 for c in corrections if c.get('block') == 'auto'),
            'textos_protegidos': len(protected)
        }
        
        # Por tipo de erro
        by_error_type = {}
        for corr in corrections:
            err_type = corr.get('error_type', 'outros')
            by_error_type[err_type] = by_error_type.get(err_type, 0) + 1
        
        report = {
            'resumo': stats,
            'por_tipo_erro': by_error_type,
            'todas_correcoes': corrections,
            'textos_protegidos': [
                {'index': p['index'], 'location': p['location'], 'preview': p['current_text'][:50]}
                for p in protected
            ]
        }
        
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        self.logger.info(f"Relatório salvo: {report_path}")