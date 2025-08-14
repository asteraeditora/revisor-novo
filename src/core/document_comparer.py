import os
import json
import shutil
import logging
from typing import List, Dict
from docx import Document
from docx.shared import RGBColor
import difflib

class DocumentComparer:
    """Compara documentos detectando ABSOLUTAMENTE TODAS as mudanças"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def compare_documents(self, original_path: str, revised_path: str, 
                         output_path: str, log_path: str = None) -> str:
        """Compara documentos detectando CADA MÍNIMA MUDANÇA"""
        
        # SEMPRE usa comparação ULTRA detalhada
        return self.create_ultra_comparison(original_path, revised_path, output_path)
    
    def create_ultra_comparison(self, original_path: str, revised_path: str, output_path: str) -> str:
        """Cria comparação detectando TUDO - cada vírgula, ponto, espaço"""
        try:
            self.logger.info("=== COMPARAÇÃO ULTRA DETALHADA ===")
            
            # 1. COPIA o documento revisado
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            shutil.copy2(revised_path, output_path)
            
            # 2. Abre os três documentos
            original_doc = Document(original_path)
            revised_doc = Document(revised_path)
            comparison_doc = Document(output_path)
            
            # 3. ANÁLISE COMPLETA - parágrafo por parágrafo
            all_changes = []
            para_num = 0
            total_paragraphs = max(len(original_doc.paragraphs), len(revised_doc.paragraphs))
            
            self.logger.info(f"Comparando {total_paragraphs} parágrafos...")
            
            # Processa TODOS os parágrafos (mesmo se um doc tem mais que o outro)
            for i in range(total_paragraphs):
                # Pega o parágrafo de cada doc (ou texto vazio se não existir)
                orig_text = original_doc.paragraphs[i].text if i < len(original_doc.paragraphs) else ""
                rev_text = revised_doc.paragraphs[i].text if i < len(revised_doc.paragraphs) else ""
                comp_para = comparison_doc.paragraphs[i] if i < len(comparison_doc.paragraphs) else None
                
                # Conta parágrafos não vazios
                if orig_text.strip() or rev_text.strip():
                    para_num += 1
                    
                    # COMPARA CARACTERE POR CARACTERE
                    if orig_text != rev_text:
                        self.logger.info(f"Mudança no parágrafo {para_num}:")
                        self.logger.info(f"  Original: '{orig_text}'")
                        self.logger.info(f"  Revisado: '{rev_text}'")
                        
                        # Analisa a diferença em DETALHE
                        changes = self._analyze_all_changes(orig_text, rev_text)
                        
                        for change in changes:
                            change.update({
                                'paragraph_number': para_num,
                                'location': f'Parágrafo {para_num}',
                                'page': (para_num // 3) + 1,
                                'original_full': orig_text,
                                'revised_full': rev_text
                            })
                            all_changes.append(change)
                        
                        # MARCA no documento
                        if comp_para:
                            self._mark_all_changes(comp_para, orig_text, rev_text)
            
            # 4. ANÁLISE COMPLETA DAS TABELAS
            self.logger.info("Analisando tabelas...")
            table_changes = self._compare_all_tables(original_doc, revised_doc, comparison_doc, para_num)
            all_changes.extend(table_changes)
            
            # 5. Adiciona sumário COMPLETO no início
            self._add_complete_summary(comparison_doc, all_changes)
            
            # 6. Salva
            comparison_doc.save(output_path)
            
            self.logger.info(f"=== COMPARAÇÃO CONCLUÍDA ===")
            self.logger.info(f"Total de {len(all_changes)} mudanças detectadas")
            self.logger.info(f"Documento salvo: {output_path}")
            
            # 7. Salva relatório JSON detalhado
            report_path = output_path.replace('.docx', '_detalhado.json')
            with open(report_path, 'w', encoding='utf-8') as f:
                json.dump({
                    'total_mudancas': len(all_changes),
                    'todas_mudancas': all_changes
                }, f, ensure_ascii=False, indent=2)
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Erro na comparação: {str(e)}")
            raise
    
    def _analyze_all_changes(self, original: str, revised: str) -> List[Dict]:
        """Analisa TODAS as mudanças entre dois textos"""
        changes = []
        
        # 1. Verifica mudanças simples primeiro
        if revised == original + '.':
            changes.append({
                'type': 'pontuação',
                'error': '[SEM PONTO FINAL]',
                'correction': '.',
                'description': 'Adicionado ponto final'
            })
            return changes
        
        if revised == original + ',':
            changes.append({
                'type': 'pontuação',
                'error': '[SEM VÍRGULA]',
                'correction': ',',
                'description': 'Adicionada vírgula'
            })
            return changes
        
        # 2. Análise CARACTERE por CARACTERE
        s = difflib.SequenceMatcher(None, original, revised)
        
        for tag, i1, i2, j1, j2 in s.get_opcodes():
            if tag == 'replace':
                changes.append({
                    'type': 'substituição',
                    'error': original[i1:i2],
                    'correction': revised[j1:j2],
                    'position': i1,
                    'description': f'"{original[i1:i2]}" → "{revised[j1:j2]}"'
                })
            elif tag == 'delete':
                changes.append({
                    'type': 'remoção',
                    'error': original[i1:i2],
                    'correction': '',
                    'position': i1,
                    'description': f'Removido: "{original[i1:i2]}"'
                })
            elif tag == 'insert':
                changes.append({
                    'type': 'adição',
                    'error': '',
                    'correction': revised[j1:j2],
                    'position': i1,
                    'description': f'Adicionado: "{revised[j1:j2]}"'
                })
        
        # Se não encontrou mudanças específicas mas os textos são diferentes
        if not changes and original != revised:
            changes.append({
                'type': 'mudança geral',
                'error': original,
                'correction': revised,
                'description': 'Texto completamente alterado'
            })
        
        return changes
    
    def _mark_all_changes(self, paragraph, original_text: str, revised_text: str):
        """Marca TODAS as mudanças visualmente"""
        
        # Limpa parágrafo
        for run in paragraph.runs:
            run.text = ""
        
        # Análise caractere por caractere
        s = difflib.SequenceMatcher(None, original_text, revised_text)
        
        for tag, i1, i2, j1, j2 in s.get_opcodes():
            if tag == 'equal':
                # Texto igual - normal
                paragraph.add_run(original_text[i1:i2])
            
            elif tag == 'replace':
                # Mostra o que foi removido
                old_run = paragraph.add_run(original_text[i1:i2])
                old_run.font.strike = True
                old_run.font.color.rgb = RGBColor(255, 0, 0)
                old_run.font.bold = True
                
                # Mostra o que foi adicionado
                new_run = paragraph.add_run(revised_text[j1:j2])
                new_run.font.color.rgb = RGBColor(0, 255, 0)
                new_run.font.underline = True
                new_run.font.bold = True
            
            elif tag == 'delete':
                # Texto removido
                del_run = paragraph.add_run(original_text[i1:i2])
                del_run.font.strike = True
                del_run.font.color.rgb = RGBColor(255, 0, 0)
                del_run.font.bold = True
            
            elif tag == 'insert':
                # Texto adicionado
                add_run = paragraph.add_run(revised_text[j1:j2])
                add_run.font.color.rgb = RGBColor(0, 255, 0)
                add_run.font.underline = True
                add_run.font.bold = True
                add_run.font.size = add_run.font.size  # Mantém tamanho
    
    def _compare_all_tables(self, original_doc, revised_doc, comparison_doc, para_offset):
        """Compara TODAS as células de TODAS as tabelas"""
        table_changes = []
        
        max_tables = max(len(original_doc.tables), len(revised_doc.tables))
        
        for t_idx in range(max_tables):
            if t_idx >= len(original_doc.tables) or t_idx >= len(revised_doc.tables):
                # Uma tabela foi adicionada ou removida
                continue
            
            orig_table = original_doc.tables[t_idx]
            rev_table = revised_doc.tables[t_idx]
            comp_table = comparison_doc.tables[t_idx] if t_idx < len(comparison_doc.tables) else None
            
            # Compara cada célula
            for r_idx in range(max(len(orig_table.rows), len(rev_table.rows))):
                if r_idx >= len(orig_table.rows) or r_idx >= len(rev_table.rows):
                    continue
                
                for c_idx in range(max(len(orig_table.rows[r_idx].cells), len(rev_table.rows[r_idx].cells))):
                    if c_idx >= len(orig_table.rows[r_idx].cells) or c_idx >= len(rev_table.rows[r_idx].cells):
                        continue
                    
                    orig_cell = orig_table.rows[r_idx].cells[c_idx]
                    rev_cell = rev_table.rows[r_idx].cells[c_idx]
                    comp_cell = comp_table.rows[r_idx].cells[c_idx] if comp_table else None
                    
                    # Compara cada parágrafo na célula
                    for p_idx in range(max(len(orig_cell.paragraphs), len(rev_cell.paragraphs))):
                        orig_text = orig_cell.paragraphs[p_idx].text if p_idx < len(orig_cell.paragraphs) else ""
                        rev_text = rev_cell.paragraphs[p_idx].text if p_idx < len(rev_cell.paragraphs) else ""
                        
                        if orig_text != rev_text and (orig_text.strip() or rev_text.strip()):
                            # Analisa mudanças
                            changes = self._analyze_all_changes(orig_text, rev_text)
                            
                            for change in changes:
                                change.update({
                                    'location': f'Tabela {t_idx+1}, Célula ({r_idx+1},{c_idx+1})',
                                    'page': (para_offset // 3) + 1,
                                    'original_full': orig_text,
                                    'revised_full': rev_text
                                })
                                table_changes.append(change)
                            
                            # Marca mudanças
                            if comp_cell and p_idx < len(comp_cell.paragraphs):
                                self._mark_all_changes(comp_cell.paragraphs[p_idx], orig_text, rev_text)
        
        return table_changes
    
    def _add_complete_summary(self, doc, all_changes):
        """Adiciona sumário ULTRA COMPLETO no início"""
        
        # Move para o início
        first_para = doc.paragraphs[0] if doc.paragraphs else None
        
        # TÍTULO
        title = doc.add_paragraph()
        title_run = title.add_run('RELATÓRIO ULTRA DETALHADO DE TODAS AS MUDANÇAS')
        title_run.bold = True
        title_run.font.size = 18
        if first_para:
            title._element.addprevious(first_para._element)
        
        # ESTATÍSTICAS
        stats = doc.add_paragraph()
        stats.add_run(f'\nTOTAL DE MUDANÇAS DETECTADAS: {len(all_changes)}\n\n')
        
        # Por tipo
        by_type = {}
        for change in all_changes:
            t = change.get('type', 'outro')
            by_type[t] = by_type.get(t, 0) + 1
        
        stats.add_run('POR TIPO:\n')
        for tipo, count in sorted(by_type.items(), key=lambda x: x[1], reverse=True):
            stats.add_run(f'• {tipo.upper()}: {count}\n')
        
        if first_para:
            stats._element.addprevious(first_para._element)
        
        # LEGENDA
        legend = doc.add_paragraph()
        legend.add_run('\nLEGENDA: ')
        
        rem = legend.add_run('texto removido')
        rem.font.strike = True
        rem.font.color.rgb = RGBColor(255, 0, 0)
        rem.font.bold = True
        
        legend.add_run(' | ')
        
        add = legend.add_run('texto adicionado')
        add.font.color.rgb = RGBColor(0, 255, 0)
        add.font.underline = True
        add.font.bold = True
        
        if first_para:
            legend._element.addprevious(first_para._element)
        
        # LISTA COMPLETA DE MUDANÇAS
        div = doc.add_paragraph()
        div.add_run('\n' + '='*100 + '\n')
        if first_para:
            div._element.addprevious(first_para._element)
        
        list_title = doc.add_paragraph()
        list_run = list_title.add_run('LISTA COMPLETA DE TODAS AS MUDANÇAS (cada vírgula, ponto, espaço)')
        list_run.bold = True
        list_run.font.size = 14
        if first_para:
            list_title._element.addprevious(first_para._element)
        
        # Lista cada mudança
        for i, change in enumerate(all_changes, 1):
            change_para = doc.add_paragraph()
            
            # Número e localização
            change_para.add_run(f'\n{i}. {change["location"]}: ').bold = True
            
            # Tipo
            type_run = change_para.add_run(f'[{change["type"].upper()}] ')
            type_run.font.color.rgb = RGBColor(0, 0, 139)
            
            # Descrição
            change_para.add_run(change['description'])
            
            # Contexto (primeiros 50 chars)
            if len(change.get('original_full', '')) > 50:
                context = change['original_full'][:50] + '...'
                change_para.add_run(f'\n   Contexto: "{context}"')
                change_para.runs[-1].font.size = 10
                change_para.runs[-1].font.color.rgb = RGBColor(128, 128, 128)
            
            if first_para:
                change_para._element.addprevious(first_para._element)
        
        # Linha final
        final = doc.add_paragraph()
        final.add_run('\n' + '='*100)
        final.add_run('\nDOCUMENTO REVISADO COM TODAS AS MARCAÇÕES:\n\n')
        if first_para:
            final._element.addprevious(first_para._element)