import os
import time
import logging
import json
import shutil
import re
from typing import List, Dict, Tuple
from docx import Document
from .modular_prompt_system import ModularPromptSystem, PromptModule
from ..utils.word_utils import WordDocumentHandler
import concurrent.futures
from threading import Lock

class SmartDocumentProcessor:
    """Processador inteligente com segmentação e prompts modulares"""
    
    def __init__(self, api_client, mode: str = "editorial"):
        self.api_client = api_client
        self.mode = mode
        self.prompt_system = ModularPromptSystem()
        self.logger = logging.getLogger(__name__)
        
        # CONFIGURAÇÕES PARA GPT-4 COM JANELA GRANDE
        self.target_paragraphs_per_block = 50  # Blocos MUITO maiores com GPT-4
        self.max_paragraphs_per_block = 100    # Aproveita janela de 128k tokens
        self.min_paragraphs_per_block = 20     # Mínimo maior também
        
        # Para processamento paralelo - reduzido para evitar sobrecarga
        self.max_workers = 2  # Apenas 2 threads simultâneas para evitar erro do servidor
        self.corrections_lock = Lock()
    
    def process_document(self, input_path: str, output_path: str, callback=None):
        """Processa documento com velocidade máxima"""
        try:
            start_time = time.time()
            self.logger.info(f"**INICIANDO PROCESSAMENTO RÁPIDO**")
            
            # 1. Copia o arquivo
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            shutil.copy2(input_path, output_path)
            
            # 2. Abre o documento
            doc = Document(output_path)
            
            # 3. Extrai conteúdo
            all_paragraphs = self._extract_all_content_fast(doc)
            total_paragraphs = len(all_paragraphs)
            self.logger.info(f"Total de parágrafos: **{total_paragraphs}**")
            
            # 4. Cria blocos PEQUENOS para processamento rápido
            blocks = self._create_fast_blocks(all_paragraphs)
            total_blocks = len(blocks)
            self.logger.info(f"Dividido em **{total_blocks} blocos pequenos** para processamento paralelo")
            
            # Estimativa de tempo para GPT-4.1
            time_per_block = 15.0  # ~15s por bloco grande com GPT-4.1 (mais realista)
            estimated_time = (total_blocks * time_per_block) / self.max_workers
            self.logger.info(f"Tempo estimado: **{estimated_time/60:.1f} minutos** (GPT-4.1 com blocos grandes)")
            
            # 5. Prepara todos os prompts de uma vez (mais eficiente)
            all_prompts = []
            
            # Usa os módulos definidos pelo modo escolhido
            modules = self.prompt_system.get_prompt_sequence(self.mode)
            self.logger.info(f"Usando {len(modules)} módulos do modo '{self.mode}'")
            
            for block_idx, block in enumerate(blocks):
                block_text = self._prepare_block_text_fast(block)
                prompt = self.prompt_system.build_prompt(modules, block_text)
                all_prompts.append((prompt, block_idx, block))
            
            # 6. Processa em PARALELO para máxima velocidade
            all_corrections = []
            
            with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # Submete todos os trabalhos
                futures = []
                for prompt, block_idx, block in all_prompts:
                    future = executor.submit(self._process_single_block_fast, prompt, block_idx, block)
                    futures.append(future)
                
                # Coleta resultados conforme ficam prontos
                completed = 0
                for future in concurrent.futures.as_completed(futures):
                    completed += 1
                    if callback:
                        callback(completed, total_blocks, f"Processando {completed}/{total_blocks}")
                    
                    try:
                        corrections = future.result()
                        if corrections:
                            all_corrections.extend(corrections)
                    except Exception as e:
                        self.logger.error(f"Erro em bloco: {str(e)}")
            
            # 7. Aplica correções
            if all_corrections:
                self._apply_corrections_fast(doc, all_corrections)
                doc.save(output_path)
                self.logger.info(f"Documento salvo com {len(all_corrections)} correções")
            
            # Tempo total
            total_time = time.time() - start_time
            self.logger.info(f"**PROCESSAMENTO CONCLUÍDO EM {total_time:.1f} SEGUNDOS**")
            self.logger.info(f"Velocidade: {total_paragraphs/total_time:.1f} parágrafos/segundo")
            
            # Salva relatório organizado por páginas
            self._save_page_report(output_path, all_corrections)
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Erro no processamento: {str(e)}")
            raise
    
    def _extract_all_content_fast(self, doc: Document) -> List[Dict]:
        """Extração com informação de página e detecção de conteúdo protegido"""
        all_content = []
        para_num = 0
        
        # Estima páginas (aproximadamente 30 parágrafos por página)
        paragraphs_per_page = 30
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                para_num += 1
                # Calcula página estimada
                estimated_page = (i // paragraphs_per_page) + 1
                
                # Detecta se é conteúdo protegido
                text = para.text.strip()
                is_protected = False
                
                # Verifica se é citação/referência
                if any(marker in text for marker in ['Fonte:', 'Referência:', 'Extraído de:', 
                                                     'Adaptado de:', 'Wikipedia', '(Fonte', 
                                                     'Disponível em:', 'Acesso em:']):
                    is_protected = True
                
                # Verifica se o próximo parágrafo é uma fonte
                if i < len(doc.paragraphs) - 1:
                    next_text = doc.paragraphs[i + 1].text.strip()
                    if any(marker in next_text for marker in ['Fonte:', 'Referência:', 'Wikipedia']):
                        is_protected = True
                
                all_content.append({
                    'index': para_num,
                    'doc_index': i,
                    'text': para.text,
                    'paragraph_obj': para,
                    'type': 'paragraph',
                    'page': estimated_page,
                    'location_text': f"Página {estimated_page}, Parágrafo {para_num}",
                    'is_protected': is_protected
                })
        
        return all_content
    
    def _create_fast_blocks(self, paragraphs: List[Dict]) -> List[List[Dict]]:
        """Cria blocos PEQUENOS para processamento rápido com proteção inteligente"""
        
        # Primeiro, marca parágrafos que pertencem a textos com referência/assinatura
        for i in range(len(paragraphs)):
            # Verifica se é assinatura ou referência
            text = paragraphs[i]['text'].strip()
            
            # Padrões de assinatura/referência
            is_signature = any(marker in text for marker in [
                'Atenciosamente', 'Cordialmente', 'Um abraço', 'Abraços',
                'Saudações', 'Respeitosamente', 'Grato', 'Obrigado'
            ])
            
            is_reference = any(marker in text for marker in [
                'Disponível em:', 'Acesso em:', 'Fonte:', 'Referência:',
                'Adaptado', 'Extraído de:', 'Retirado de:', 'https://', 'http://', 'www.'
            ])
            
            # Se encontrou assinatura ou referência, marca TODO o texto acima como protegido
            if is_signature or is_reference:
                # Procura o início do texto (título ou após questão)
                start_idx = i
                
                # Volta procurando o início do texto
                for j in range(i-1, -1, -1):
                    para_text = paragraphs[j]['text'].strip()
                    
                    # Para se encontrar:
                    # - Questão numerada
                    # - Parágrafo vazio
                    # - Outro texto com referência
                    if (not para_text or 
                        re.match(r'^\d+[\.\)]', para_text) or
                        any(m in para_text for m in ['Disponível em:', 'Fonte:'])):
                        start_idx = j + 1
                        break
                    
                    # Se encontrou possível título (linha curta sem ponto final)
                    if len(para_text) < 100 and not para_text.endswith('.'):
                        start_idx = j
                
                # Marca todos os parágrafos do texto como protegidos
                for k in range(start_idx, i + 1):
                    if k < len(paragraphs):
                        paragraphs[k]['is_protected'] = True
                        paragraphs[k]['protection_reason'] = 'texto_com_referencia'
        
        # Agora cria os blocos, respeitando proteções
        blocks = []
        current_block = []
        
        for para in paragraphs:
            # Se é protegido, não inclui no bloco de processamento
            if not para.get('is_protected', False):
                current_block.append(para)
            
            # Quando atingir o tamanho do bloco ou encontrar um protegido, fecha o bloco
            if len(current_block) >= self.target_paragraphs_per_block or para.get('is_protected', False):
                if current_block:
                    blocks.append(current_block)
                    current_block = []
        
        # Adiciona último bloco se houver
        if current_block:
            blocks.append(current_block)
        
        return blocks
    
    def _prepare_block_text_fast(self, block: List[Dict]) -> str:
        """Prepara texto com informação de localização e identificação de questões"""
        lines = []
        
        # Detecta se o bloco contém questões com gabarito
        has_questions = False
        gabarito_info = {}
        
        # Primeiro, procura por gabaritos no bloco
        for i, para in enumerate(block):
            text = para['text'].strip()
            # Detecta gabarito (várias formas comuns)
            if any(marker in text.lower() for marker in ['gabarito:', 'resposta:', 'alternativa correta:', 'letra correta:']):
                has_questions = True
                # Extrai a letra do gabarito
                import re
                match = re.search(r'[:\s]([A-Ea-e])[)\s\.]', text)
                if match:
                    correct_letter = match.group(1).upper()
                    # Procura a questão acima do gabarito
                    for j in range(i-1, max(i-10, -1), -1):  # Procura até 10 parágrafos acima
                        if j >= 0 and re.match(r'^\d+[\.\)]\s*', block[j]['text']):
                            gabarito_info[j] = correct_letter
                            break
        
        # Prepara o texto com marcações especiais
        for i, para in enumerate(block):
            text = para['text']
            para_num = para['index']
            page = para.get('page', '?')
            
            # Se é uma alternativa de questão com gabarito conhecido
            if has_questions and re.match(r'^[a-eA-E][\)\.]', text):
                alt_letter = text[0].upper()
                
                # Procura o gabarito desta questão
                question_idx = None
                for j in range(i-1, max(i-6, -1), -1):
                    if j in gabarito_info:
                        question_idx = j
                        break
                
                if question_idx is not None and question_idx in gabarito_info:
                    correct_letter = gabarito_info[question_idx]
                    if alt_letter == correct_letter:
                        lines.append(f"[P{para_num}][PÁG{page}][ALTERNATIVA_CORRETA] {text}")
                    else:
                        lines.append(f"[P{para_num}][PÁG{page}][ALTERNATIVA_INCORRETA] {text}")
                else:
                    lines.append(f"[P{para_num}][PÁG{page}] {text}")
            else:
                lines.append(f"[P{para_num}][PÁG{page}] {text}")
        
        return "\n".join(lines)
    
    def _process_single_block_fast(self, prompt: str, block_idx: int, block: List[Dict]) -> List[Dict]:
        """Processa um único bloco rapidamente"""
        try:
            # Chama API
            corrections = self.api_client.identify_errors_precise(prompt, block_idx)
            
            # Mapeia correções para índices globais
            if corrections:
                for corr in corrections:
                    para_num_in_block = corr.get('paragraph', 0)
                    if 0 < para_num_in_block <= len(block):
                        global_para_index = block[para_num_in_block - 1]['index']
                        corr['paragraph'] = global_para_index
            
            return corrections
            
        except Exception as e:
            self.logger.error(f"Erro no bloco {block_idx}: {str(e)}")
            return []
    
    def _apply_corrections_fast(self, doc: Document, corrections: List[Dict]):
        """Aplicação otimizada com relatório por página"""
        # Agrupa por página primeiro
        corrections_by_page = {}
        
        # Mapeia parágrafos com páginas
        para_map = {}
        para_counter = 0
        paragraphs_per_page = 30
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                para_counter += 1
                estimated_page = (i // paragraphs_per_page) + 1
                para_map[para_counter] = {
                    'paragraph': para,
                    'page': estimated_page
                }
        
        # Organiza correções por página
        for corr in corrections:
            para_num = corr.get('paragraph', 0)
            if para_num in para_map:
                page = para_map[para_num]['page']
                if page not in corrections_by_page:
                    corrections_by_page[page] = []
                corr['page'] = page
                corr['location'] = f"Página {page}, Parágrafo {para_num}"
                corrections_by_page[page].append(corr)
        
        # Log de correções por página
        self.logger.info("**CORREÇÕES POR PÁGINA:**")
        for page in sorted(corrections_by_page.keys()):
            self.logger.info(f"  Página {page}: {len(corrections_by_page[page])} correções")
        
        # Aplica correções
        applied_count = 0
        for corr in corrections:
            para_num = corr.get('paragraph', 0)
            if para_num in para_map:
                paragraph = para_map[para_num]['paragraph']
                current_text = paragraph.text
                
                error = corr.get('error', '')
                correction = corr.get('correction', '')
                
                if error and correction and error in current_text:
                    current_text = current_text.replace(error, correction, 1)
                    paragraph.text = current_text
        
        self.logger.info(f"Total de correções aplicadas: {applied_count}")
    
    def _save_page_report(self, output_path: str, corrections: List[Dict]):
        """Salva relatório organizado por páginas"""
        # Organiza por página
        by_page = {}
        for corr in corrections:
            page = corr.get('page', 0)
            if page not in by_page:
                by_page[page] = []
            by_page[page].append({
                'paragrafo': corr.get('paragraph'),
                'erro': corr.get('error'),
                'correcao': corr.get('correction'),
                'tipo': corr.get('type', corr.get('module', 'outro')),
                'localizacao': corr.get('location', '')
            })
        
        # Cria relatório
        report = {
            'total_correções': len(corrections),
            'total_páginas_com_correções': len(by_page),
            'correções_por_página': {}
        }
        
        # Adiciona correções ordenadas por página
        for page in sorted(by_page.keys()):
            report['correções_por_página'][f'página_{page}'] = {
                'total': len(by_page[page]),
                'correções': by_page[page]
            }
        
        # Salva JSON
        report_path = output_path.replace('.docx', '_relatorio_paginas.json')
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        self.logger.info(f"Relatório salvo em: {report_path}")